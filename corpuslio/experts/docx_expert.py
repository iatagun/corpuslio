"""DOCX Expert Module for OCRchestra.

Extracts text from Microsoft Word (.docx) files using python-docx.
Preserves paragraph structure and can optionally extract tables.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from .base import ExpertBase

logger = logging.getLogger(__name__)

# Try to import docx
try:
    from docx import Document
    from docx.table import Table
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None


class DOCXExpert(ExpertBase):
    """Expert for extracting text from DOCX documents."""

    @staticmethod
    def is_available() -> bool:
        """Check if python-docx is installed."""
        return HAS_DOCX

    def extract_text(
        self,
        docx_path: Union[str, Path],
        include_tables: bool = True,
        preserve_structure: bool = True,
    ) -> Dict[str, Any]:
        """Extract text from DOCX file.

        Args:
            docx_path: Path to DOCX file
            include_tables: Whether to extract table contents
            preserve_structure: Whether to preserve paragraph breaks

        Returns:
            Dict with extracted text and metadata
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx yüklü değil. Lütfen: pip install python-docx",
                "text": "",
            }

        docx_path = Path(docx_path)
        if not docx_path.exists():
            return {
                "success": False,
                "error": f"DOCX dosyası bulunamadı: {docx_path}",
                "text": "",
            }

        try:
            doc = Document(docx_path)
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX açılamadı: {e}",
                "text": "",
            }

        paragraphs = []
        tables_data = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract tables if requested
        if include_tables:
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table(table)
                if table_text:
                    tables_data.append({
                        "index": table_idx + 1,
                        "text": table_text,
                    })

        # Combine text
        if preserve_structure:
            full_text = "\n\n".join(paragraphs)
        else:
            full_text = " ".join(paragraphs)

        # Add tables to text
        if tables_data:
            table_texts = [f"[Tablo {t['index']}]\n{t['text']}" for t in tables_data]
            full_text += "\n\n" + "\n\n".join(table_texts)

        return {
            "success": True,
            "file": str(docx_path),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables_data),
            "text": full_text,
            "paragraphs": paragraphs,
            "tables": tables_data,
        }

    def _extract_table(self, table) -> str:
        """Extract text from a table.

        Args:
            table: python-docx Table object

        Returns:
            Formatted table text
        """
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        return "\n".join(rows_text)

    @staticmethod
    def execute(input_json: Dict[str, Any]) -> Dict[str, Any]:
        """Execute DOCX extraction (stateless expert interface).

        Args:
            input_json: Dict with 'file_path' and optional 'include_tables'

        Returns:
            Extraction result dict
        """
        file_path = input_json.get("file_path")
        if not file_path:
            return {"success": False, "error": "file_path gerekli"}

        include_tables = input_json.get("include_tables", True)
        preserve_structure = input_json.get("preserve_structure", True)

        expert = DOCXExpert()
        return expert.extract_text(
            file_path,
            include_tables=include_tables,
            preserve_structure=preserve_structure,
        )


def execute(input_json: Dict[str, Any]) -> Dict[str, Any]:
    """Module-level execute function."""
    return DOCXExpert.execute(input_json)
